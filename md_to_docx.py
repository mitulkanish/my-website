import docx
import re

def markdown_to_docx(md_filepath, docx_filepath):
    doc = docx.Document()
    
    with open(md_filepath, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    in_code_block = False
    
    for line in lines:
        line_clean = line.strip()
        
        if line_clean.startswith('```'):
            in_code_block = not in_code_block
            continue
            
        if in_code_block:
            p = doc.add_paragraph(line)
            continue
            
        if line_clean == '---':
            doc.add_page_break()
            continue
            
        if not line_clean:
            continue
            
        if line_clean.startswith('# '):
            doc.add_heading(line_clean[2:], level=1)
        elif line_clean.startswith('## '):
            doc.add_heading(line_clean[3:], level=2)
        elif line_clean.startswith('### '):
            doc.add_heading(line_clean[4:], level=3)
        else:
            p = doc.add_paragraph()
            
            if line_clean.startswith('- ') or line_clean.startswith('* '):
                p.style = 'List Bullet'
                line = line_clean[2:]
            elif re.match(r'^\d+\.\s', line_clean):
                 p.style = 'List Number'
                 line = re.sub(r'^\d+\.\s', '', line_clean)
                 
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    subparts = re.split(r'(`.*?`)', part)
                    for sp in subparts:
                        if sp.startswith('`') and sp.endswith('`'):
                            run = p.add_run(sp[1:-1])
                        else:
                            p.add_run(sp)


    doc.save(docx_filepath)
    print("Done")

if __name__ == "__main__":
    markdown_to_docx("README.md", "INTERVENIX_Documentation.docx")
